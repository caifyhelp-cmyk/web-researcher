# -*- coding: utf-8 -*-
"""
MAESTRO v1.1 — 현존 최강 AI 오케스트레이터

조경일 뇌 에이전트 × GPT-4o × Claude Code × DeepSeek × Grok
Claude Code 동일 Tool Suite + ReAct 에이전트 루프

아키텍처:
  - GPT-4o      : 메인 오케스트레이터 (도구 호출, 흐름 제어)
  - Claude Code : 실제 Claude Code CLI 호출 (코딩/파일작업 최강)
  - DeepSeek    : 복잡한 추론, 수학, 단계적 분석
  - Claude API  : 긴 문서, 정밀 분석, 세밀한 글쓰기
  - Grok        : 실시간 정보, 최신 이슈
  - 뇌 에이전트 : 마케팅/전략 판단 (조경일 인지 패턴 1,180개)
"""

import os, sys, json, re, subprocess, time
from pathlib import Path
from datetime import datetime

VERSION = "2.0.0"

# ── Rich UI ──────────────────────────────────────────────────────
from rich.console import Console
from rich.panel import Panel
from rich.markdown import Markdown
from rich.prompt import Prompt, Confirm
from rich.rule import Rule
from rich.live import Live
from rich.spinner import Spinner
from rich.text import Text

console = Console(highlight=False)

# ── API 키 ───────────────────────────────────────────────────────
try:
    import _local_keys
except ImportError:
    pass

OPENAI_KEY    = os.getenv("OPENAI_API_KEY", "")
ANTHROPIC_KEY = os.getenv("ANTHROPIC_API_KEY", "")
DEEPSEEK_KEY  = os.getenv("DEEPSEEK_API_KEY", "")
GROK_KEY      = os.getenv("GROK_API_KEY", "")
GEMINI_KEY    = os.getenv("GEMINI_API_KEY", "")
VERCEL_TOKEN  = os.getenv("VERCEL_TOKEN", "")

from openai    import OpenAI
from anthropic import Anthropic

oai       = OpenAI(api_key=OPENAI_KEY)                                          if OPENAI_KEY    else None
ant       = Anthropic(api_key=ANTHROPIC_KEY)                                    if ANTHROPIC_KEY else None
deepseek  = OpenAI(api_key=DEEPSEEK_KEY, base_url="https://api.deepseek.com")  if DEEPSEEK_KEY  else None
grok_ai   = OpenAI(api_key=GROK_KEY,     base_url="https://api.x.ai/v1")       if GROK_KEY      else None
gemini_ai = OpenAI(api_key=GEMINI_KEY,
                   base_url="https://generativelanguage.googleapis.com/v1beta/openai/") if GEMINI_KEY else None

# ── 자동 업데이트 ────────────────────────────────────────────────
_GITHUB_RAW = "https://raw.githubusercontent.com/caifyhelp-cmyk/web-researcher/master"

def _check_update():
    """GitHub version.txt 확인 후 전체 파일 자동 업데이트 및 재시작"""
    import urllib.request as _ur
    try:
        with _ur.urlopen(f"{_GITHUB_RAW}/version.txt", timeout=4) as r:
            latest = r.read().decode().strip()
        # app_local.py 의 VERSION 기준으로 비교
        try:
            import app_local as _al
            cur = getattr(_al, "VERSION", VERSION)
        except Exception:
            cur = VERSION
        if latest == cur:
            return
        print(f"  업데이트 발견: {cur} -> {latest}  다운로드 중...")
        here = Path(__file__).parent
        self_ok = False
        files = [
            ("app_local.py",          "app_local.py"),
            ("maestro.py",            "maestro.py"),
            ("web_researcher.py",     "web_researcher.py"),
            ("orchestrator.py",       "orchestrator.py"),
            ("feedback_collector.py", "feedback_collector.py"),
            ("pattern_collector.py",  "pattern_collector.py"),
            ("tools_kb.json",         "tools_kb.json"),
        ]
        for gh_name, local_name in files:
            try:
                import time as _t
                tmp, _ = _ur.urlretrieve(f"{_GITHUB_RAW}/{gh_name}?t={int(_t.time())}")
                dest = here / local_name
                raw = open(tmp, "rb").read()
                open(dest, "wb").write(raw)
                if local_name == "app_local.py":
                    import re as _re
                    txt = raw.decode(errors="replace")
                    m = _re.search(r'VERSION\s*=\s*["\']([^"\']+)["\']', txt)
                    if m and m.group(1) == latest:
                        self_ok = True
            except Exception:
                pass
        if not self_ok:
            print(f"  업데이트 파일 검증 실패 (VERSION 불일치). 재시작 건너뜀.")
            return
        print("  업데이트 완료. 재시작합니다...\n")
        time.sleep(1)
        subprocess.Popen([sys.executable, os.path.abspath(__file__)] + sys.argv[1:])
        os._exit(0)
    except Exception:
        pass

# ── 오케스트레이터 ────────────────────────────────────────────────
try:
    import orchestrator as orch
    _ORCH = True
except Exception:
    _ORCH = False

# ── 웹 리서치 파이프라인 (app_local.py 통합) ─────────────────────
try:
    import app_local as _rl
    _RL = True
except Exception:
    _RL = False

# ── 집단지성 파이프라인 ───────────────────────────────────────────
try:
    import pattern_collector as _pc
    _PC = True
    # 시작 시 GitHub에서 최신 지식베이스 동기화
    try:
        _pc.pull_kb_from_github()
    except Exception:
        pass
except Exception:
    _PC = False

# ── 툴 지식 DB ────────────────────────────────────────────────────
_TOOLS_KB_PATH = Path(__file__).parent / "tools_kb.json"

def _load_tools_kb() -> list:
    try:
        return json.loads(_TOOLS_KB_PATH.read_text(encoding="utf-8")).get("tools", [])
    except Exception:
        return []

# 마지막 리서치 결과 캐시 (save_research 도구에서 재사용)
_last_research: dict = {}

# ── Selenium 싱글턴 드라이버 ──────────────────────────────────────
_selenium_driver = None
_selenium_ok: bool | None = None   # None=미시도, True=성공, False=불가

def _get_selenium_driver():
    """
    web_researcher.make_driver()를 재사용해 싱글턴 Selenium 드라이버 반환.
    Chrome 없거나 selenium 미설치 시 None 반환 (requests 모드로 폴백).
    """
    global _selenium_driver, _selenium_ok
    if _selenium_ok is False:
        return None
    if _selenium_driver is not None:
        try:
            _ = _selenium_driver.current_url   # 살아있는지 확인
            return _selenium_driver
        except Exception:
            _selenium_driver = None

    try:
        from web_researcher import make_driver
        _selenium_driver = make_driver()
        _selenium_ok = _selenium_driver is not None
        if _selenium_ok:
            console.print("  [dim][Selenium] Chrome 드라이버 초기화 완료[/dim]")
        return _selenium_driver
    except Exception:
        _selenium_ok = False
        return None

# ── 쿼리 복잡도 티어 판단 ─────────────────────────────────────────
def _judge_query_tier(text: str) -> str:
    """
    사용자 입력의 복잡도를 판단해 'medium' / 'complex' 반환.
    오케스트레이터(gpt-4.1)는 항상 실행됨 — 바이패스 없음.

    medium  → gpt-4.1 표준 ReAct 루프 (오케스트레이터가 적절한 specialist 선택)
    complex → gpt-4.1 루프 + 첫 도구로 ask_ensemble 또는 ask_claude_code 강제
    """
    t = text.lower()
    complex_kw = [
        # 전략·분석
        "조사", "분석", "보고서", "비교", "시장", "경쟁사", "전략", "마케팅",
        "제안서", "기획", "리포트", "인사이트", "트렌드",
        # 코딩·개발
        "코드", "코딩", "개발", "구현", "작성해", "만들어", "짜줘", "수정",
        "리팩터", "버그", "앱", "웹사이트", "함수", "클래스", "api", "서버",
        # 리서치
        "research", "analyze", "plan", "앙상블", "종합",
    ]
    if any(k in t for k in complex_kw):
        return "complex"
    return "medium"


# ── 앙상블 병렬 호출 ──────────────────────────────────────────────
def _tool_ask_ensemble(prompt: str, category: str = "strategy_insight") -> str:
    """
    claude + gemini + deepseek 를 병렬 호출 → gpt-4.1 이 종합.
    최고 품질이 필요한 전략·분석·문서 작업에 사용.
    """
    import concurrent.futures

    models = [
        ("claude",   lambda p: _tool_ask_specialist("claude",   p, category)),
        ("gemini",   lambda p: _tool_ask_specialist("gemini",   p, category)),
        ("deepseek", lambda p: _tool_ask_specialist("deepseek", p, category)),
    ]

    results: dict[str, str] = {}

    def _call(name, fn):
        try:
            return name, fn(prompt)
        except Exception as e:
            return name, f"[{name} 오류: {e}]"

    console.print("  [ENSEMBLE] [dim]Claude · Gemini · DeepSeek 병렬 호출 중...[/dim]")
    with concurrent.futures.ThreadPoolExecutor(max_workers=3) as ex:
        futs = {ex.submit(_call, name, fn): name for name, fn in models}
        for fut in concurrent.futures.as_completed(futs):
            name, res = fut.result()
            results[name] = res
            console.print(f"    [dim]✓ {name} 완료[/dim]")

    # 유효한 응답만 수집
    valid = {k: v for k, v in results.items() if not v.startswith("[") or "오류" not in v}
    if not valid:
        return "[앙상블: 모든 모델 실패]\n" + "\n".join(f"{k}: {v}" for k, v in results.items())

    if len(valid) == 1:
        return list(valid.values())[0]

    # gpt-4.1 이 종합
    if not oai:
        # OpenAI 없으면 첫 번째 결과 반환
        return list(valid.values())[0]

    synthesis_prompt = (
        "다음은 같은 질문에 대한 여러 AI의 답변입니다. "
        "각 답변의 핵심을 통합해 가장 완성도 높은 한국어 답변을 작성하세요. "
        "중복 제거, 상충 시 더 구체적인 것 채택, 누락된 인사이트 보완.\n\n"
        f"[원래 질문]\n{prompt}\n\n"
    )
    for name, res in valid.items():
        synthesis_prompt += f"[{name.upper()} 답변]\n{res[:2000]}\n\n"

    try:
        console.print("  [ENSEMBLE] [dim]gpt-4.1 종합 중...[/dim]")
        r = oai.chat.completions.create(
            model="gpt-4.1",
            messages=[{"role": "user", "content": synthesis_prompt}],
            max_tokens=4000,
            temperature=0.4
        )
        return r.choices[0].message.content.strip()
    except Exception as e:
        return list(valid.values())[0] + f"\n[종합 오류: {e}]"


# ── 회의록 자동화 경로 ────────────────────────────────────────────
_MEETING_AUTO_DIR = Path(os.path.expanduser("~")) / "Desktop" / "meeting_auto"

# ── 응답 캐시 (SQLite + 임베딩) ───────────────────────────────────
import sqlite3

_CACHE_DB_PATH = Path(os.path.expanduser("~")) / ".maestro" / "response_cache.db"
_CACHE_DB_PATH.parent.mkdir(parents=True, exist_ok=True)
_CACHE_HIT_THRESHOLD  = 0.97   # 이 이상: 캐시를 참고로 주입 (GPT-4o가 검증 후 활용)
_CACHE_TTL_DAYS       = 14     # 캐시 유효 기간 (2주)

def _cache_init():
    conn = sqlite3.connect(str(_CACHE_DB_PATH))
    conn.execute("""
        CREATE TABLE IF NOT EXISTS response_cache (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            query_hash  TEXT NOT NULL,
            query_text  TEXT NOT NULL,
            embedding   TEXT NOT NULL,
            response    TEXT NOT NULL,
            created_at  TEXT NOT NULL,
            hit_count   INTEGER DEFAULT 0
        )
    """)
    conn.commit()
    conn.close()

_cache_init()

def _embed_query(text: str) -> list:
    """짧은 쿼리를 임베딩으로 변환 (text-embedding-3-small, 초저가)"""
    if not oai:
        return []
    try:
        resp = oai.embeddings.create(model="text-embedding-3-small", input=[text[:500]])
        return resp.data[0].embedding
    except Exception:
        return []

def _cache_lookup(query: str) -> tuple:
    """
    유사한 이전 응답 검색.
    반환: (cached_response: str, similarity: float)
    similarity >= _CACHE_HIT_THRESHOLD 이면 참고로 활용 가능.
    없으면 ("", 0.0)
    """
    try:
        import numpy as np
        q_vec = _embed_query(query)
        if not q_vec:
            return "", 0.0

        conn = sqlite3.connect(str(_CACHE_DB_PATH))
        cutoff_str = datetime.fromtimestamp(
            datetime.now().timestamp() - _CACHE_TTL_DAYS * 86400
        ).isoformat()
        rows = conn.execute(
            "SELECT id, query_text, embedding, response FROM response_cache WHERE created_at > ?",
            (cutoff_str,)
        ).fetchall()
        conn.close()

        if not rows:
            return "", 0.0

        q_arr = np.array(q_vec, dtype=np.float32)
        q_arr /= (np.linalg.norm(q_arr) + 1e-9)

        best_sim, best_row = 0.0, None
        for row in rows:
            try:
                emb = np.array(json.loads(row[2]), dtype=np.float32)
                emb /= (np.linalg.norm(emb) + 1e-9)
                sim = float(q_arr @ emb)
                if sim > best_sim:
                    best_sim, best_row = sim, row
            except Exception:
                continue

        if best_sim >= _CACHE_HIT_THRESHOLD and best_row:
            conn2 = sqlite3.connect(str(_CACHE_DB_PATH))
            conn2.execute("UPDATE response_cache SET hit_count = hit_count + 1 WHERE id = ?",
                          (best_row[0],))
            conn2.commit()
            conn2.close()
            return best_row[3], best_sim

    except Exception:
        pass
    return "", 0.0

def _cache_store(query: str, response: str):
    """응답을 캐시에 저장 (도구 미사용 순수 LLM 응답만)"""
    try:
        import hashlib
        q_vec = _embed_query(query)
        if not q_vec:
            return
        qhash = hashlib.md5(query.encode()).hexdigest()
        conn = sqlite3.connect(str(_CACHE_DB_PATH))
        conn.execute(
            "INSERT INTO response_cache (query_hash, query_text, embedding, response, created_at) VALUES (?,?,?,?,?)",
            (qhash, query[:500], json.dumps(q_vec), response, datetime.now().isoformat())
        )
        conn.commit()
        conn.close()
    except Exception:
        pass

# 저장 폴더
_SAVE_DIR = Path(os.path.expanduser("~")) / "Desktop" / "MAESTRO결과"
_SAVE_DIR.mkdir(parents=True, exist_ok=True)

# ── 세션 로그 ─────────────────────────────────────────────────────
_SESSIONS_DIR = Path(os.path.expanduser("~")) / ".maestro" / "sessions"
_SESSIONS_DIR.mkdir(parents=True, exist_ok=True)

# 하위 호환 — 구버전 단일 파일 마이그레이션
_LEGACY_HISTORY = Path(os.path.expanduser("~")) / ".maestro_history.json"

_current_session_id: str = ""


def _new_session_id() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def _session_path(sid: str) -> Path:
    return _SESSIONS_DIR / f"{sid}.json"


def _save_session(sid: str, history: list, first_msg: str = ""):
    """현재 세션 전체를 로컬에 저장 (무제한 로그)"""
    try:
        path = _session_path(sid)
        data: dict = {}
        if path.exists():
            try:
                data = json.loads(path.read_text(encoding="utf-8"))
            except Exception:
                data = {}
        data["session_id"]  = sid
        data["last_active"] = datetime.now().isoformat()
        data["turn_count"]  = len(history) // 2
        if not data.get("started_at"):
            data["started_at"] = datetime.now().isoformat()
        if not data.get("first_message") and first_msg:
            data["first_message"] = first_msg[:60]
        data["messages"] = history          # 전체 저장 (무제한)
        path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception:
        pass


def _load_session(sid: str) -> list:
    """세션 ID로 전체 히스토리 로드"""
    try:
        path = _session_path(sid)
        if path.exists():
            data = json.loads(path.read_text(encoding="utf-8"))
            return data.get("messages", [])
    except Exception:
        pass
    return []


def _list_sessions(n: int = 10) -> list:
    """최근 세션 목록 반환. [{sid, started_at, turn_count, first_message}, ...]"""
    sessions = []
    for p in sorted(_SESSIONS_DIR.glob("*.json"), reverse=True)[:n]:
        try:
            data = json.loads(p.read_text(encoding="utf-8"))
            sessions.append({
                "sid":           data.get("session_id", p.stem),
                "started_at":    data.get("started_at", "")[:16].replace("T", " "),
                "last_active":   data.get("last_active", "")[:16].replace("T", " "),
                "turn_count":    data.get("turn_count", 0),
                "first_message": data.get("first_message", "(내용 없음)"),
            })
        except Exception:
            continue
    return sessions

# ── 뇌 에이전트 ──────────────────────────────────────────────────
_BRAIN_URL = "https://brain-agent-v9wl.onrender.com/api/research"

import urllib.request, urllib.error

def _call_brain(situation: str) -> str:
    ping_url  = _BRAIN_URL.replace("/api/research", "/")
    api_url   = _BRAIN_URL

    # 콜드 스타트 대응: 먼저 홈 핑해서 깨우기
    try:
        urllib.request.urlopen(ping_url, timeout=5)
    except Exception:
        pass

    try:
        body = json.dumps({"situation": situation}, ensure_ascii=False).encode()
        req  = urllib.request.Request(api_url, data=body, method="POST",
                                      headers={"Content-Type": "application/json"})
        with urllib.request.urlopen(req, timeout=90) as r:   # 콜드 스타트 고려 90초
            data = json.loads(r.read())
        if data.get("ok"):
            parts = []
            if data.get("judgment"): parts.append(f"판단: {data['judgment']}")
            if data.get("action"):   parts.append(f"액션: {data['action']}")
            if data.get("reason"):   parts.append(f"근거: {data['reason']}")
            return "\n".join(parts) or "[뇌 에이전트 응답 없음]"
        err = data.get("error", "")
        return f"[뇌 에이전트 오류: {err}]"
    except Exception as e:
        return f"[뇌 에이전트 연결 실패: {e}]"


# ═══════════════════════════════════════════════════════════════
#  TOOL SUITE — Claude Code 동일 11개 도구
# ═══════════════════════════════════════════════════════════════

def _tool_read_file(path: str, offset: int = 0, limit: int = 0) -> str:
    try:
        p = Path(path).expanduser()
        if not p.exists():
            return f"[파일 없음: {path}]"
        lines = p.read_text(encoding="utf-8", errors="replace").splitlines()
        if offset:
            lines = lines[offset:]
        if limit:
            lines = lines[:limit]
        numbered = [f"{i+1+offset}\t{l}" for i, l in enumerate(lines)]
        return "\n".join(numbered)
    except Exception as e:
        return f"[읽기 오류: {e}]"


def _tool_write_file(path: str, content: str) -> str:
    try:
        p = Path(path).expanduser()
        p.parent.mkdir(parents=True, exist_ok=True)
        p.write_text(content, encoding="utf-8")
        return f"[완료] {path} 저장됨 ({len(content.splitlines())}줄)"
    except Exception as e:
        return f"[쓰기 오류: {e}]"


def _tool_edit_file(path: str, old_str: str, new_str: str) -> str:
    try:
        p = Path(path).expanduser()
        if not p.exists():
            return f"[파일 없음: {path}]"
        content = p.read_text(encoding="utf-8")
        if old_str not in content:
            return f"[매칭 실패] '{old_str[:60]}...' 를 파일에서 찾을 수 없음"
        new_content = content.replace(old_str, new_str, 1)
        p.write_text(new_content, encoding="utf-8")
        return f"[완료] {path} 수정됨"
    except Exception as e:
        return f"[수정 오류: {e}]"


_DANGEROUS = ["rm -rf", "format", "del /f", "shutdown", "drop table",
               "DROP TABLE", ":(){", "mkfs", "dd if="]

def _tool_run_bash(command: str, timeout: int = 30, _confirmed: bool = False) -> str:
    if not _confirmed:
        for danger in _DANGEROUS:
            if danger.lower() in command.lower():
                return f"[차단] 위험 명령어 감지: '{danger}'. 직접 실행하려면 확인 필요."
    try:
        result = subprocess.run(
            command, shell=True, capture_output=True,
            text=True, timeout=timeout, encoding="utf-8", errors="replace"
        )
        out = (result.stdout + result.stderr).strip()
        return out[:4000] if out else "[출력 없음]"
    except subprocess.TimeoutExpired:
        return f"[타임아웃] {timeout}초 초과"
    except Exception as e:
        return f"[실행 오류: {e}]"


def _tool_glob(pattern: str, path: str = ".") -> str:
    try:
        base = Path(path).expanduser()
        matches = sorted(base.glob(pattern))
        if not matches:
            return "[결과 없음]"
        return "\n".join(str(m) for m in matches[:100])
    except Exception as e:
        return f"[glob 오류: {e}]"


def _tool_grep(pattern: str, path: str = ".", file_pattern: str = "") -> str:
    try:
        cmd = ["grep", "-rn", "--include", file_pattern or "*", pattern, path]
        result = subprocess.run(cmd, capture_output=True, text=True,
                                encoding="utf-8", errors="replace", timeout=15)
        out = result.stdout.strip()
        if not out:
            # Windows fallback — findstr
            cmd2 = f'findstr /s /n /i "{pattern}" "{path}\\*{file_pattern}"'
            result2 = subprocess.run(cmd2, shell=True, capture_output=True,
                                     text=True, encoding="utf-8", errors="replace", timeout=15)
            out = result2.stdout.strip()
        return out[:4000] if out else "[결과 없음]"
    except Exception as e:
        return f"[grep 오류: {e}]"


def _tool_list_dir(path: str = ".") -> str:
    try:
        p = Path(path).expanduser()
        if not p.exists():
            return f"[경로 없음: {path}]"
        items = sorted(p.iterdir())
        lines = []
        for item in items[:200]:
            tag  = "D" if item.is_dir() else "F"
            size = item.stat().st_size if item.is_file() else 0
            lines.append(f"[{tag}] {item.name}  {size:,}B" if size else f"[{tag}] {item.name}")
        return "\n".join(lines) or "[비어있음]"
    except Exception as e:
        return f"[목록 오류: {e}]"


def _tool_web_search(query: str, num_results: int = 5) -> str:
    try:
        from duckduckgo_search import DDGS
        results = []
        with DDGS() as ddgs:
            for r in ddgs.text(query, max_results=num_results):
                results.append(f"[{r['title']}]\n{r['href']}\n{r['body']}\n")
        return "\n---\n".join(results) if results else "[결과 없음]"
    except Exception as e:
        return f"[검색 오류: {e}]"


def _tool_web_fetch(url: str) -> str:
    """
    URL 페이지 본문 추출.
    1차: requests + BeautifulSoup (빠름)
    2차: Selenium (JS 렌더링 필요한 동적 사이트 — 텍스트 200자 미만일 때 자동 폴백)
    """
    import requests as _req
    from bs4 import BeautifulSoup as _BS

    _HEADERS = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/124.0 Safari/537.36",
        "Accept-Language": "ko-KR,ko;q=0.9,en;q=0.8",
    }

    text = ""
    fetch_ok = False

    # ── 1차: requests + BeautifulSoup ────────────────────────────
    try:
        resp = _req.get(url, headers=_HEADERS, timeout=10, allow_redirects=True)
        resp.raise_for_status()
        soup = _BS(resp.content, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
            tag.decompose()
        text = " ".join(soup.get_text(separator=" ").split())
        fetch_ok = True
    except Exception as e:
        text = ""
        fetch_ok = False

    # ── 2차: Selenium 폴백 (JS 사이트 감지) ──────────────────────
    # 텍스트가 200자 미만이면 JS 렌더링 필요 → Selenium 시도
    if len(text) < 200:
        driver = _get_selenium_driver()
        if driver is not None:
            try:
                from web_researcher import scrape_page
                console.print(f"  [dim][Selenium] JS 렌더링 시도: {url[:50]}[/dim]")
                page = scrape_page(driver, url)
                sel_text = page.get("full_text", "")
                if len(sel_text) > len(text):
                    headings = page.get("headings", "")
                    combined = f"{headings}\n{sel_text}".strip() if headings else sel_text
                    return combined[:5000]
            except Exception as e:
                console.print(f"  [dim][Selenium] 오류: {e}[/dim]")

    if not fetch_ok and not text:
        return f"[fetch 오류: 페이지를 가져올 수 없습니다]"

    return text[:5000] if text else "[페이지 내용 없음]"


def _tool_ask_specialist(model: str, prompt: str, category: str = "") -> str:
    """전문 LLM에게 위임. model="auto"이면 오케스트레이터가 최적 모델 선택."""

    # ── auto: 오케스트레이터 동적 선택 ──────────────────────────────
    if model == "auto":
        if _ORCH and category:
            if orch.check_reevaluation_needed(category):
                pass  # 백그라운드 재평가는 main loop에서 처리
            model = orch.get_best_model(category)
        else:
            model = "gpt-4o"

    if model == "deepseek":
        if not deepseek:
            return "[DeepSeek API 키 없음]"
        try:
            r = deepseek.chat.completions.create(
                model="deepseek-reasoner",   # 내부: DeepSeek-R2 (API명 동일)
                messages=[{"role": "user", "content": prompt}],
                max_tokens=3000
            )
            return r.choices[0].message.content.strip()
        except Exception as e:
            return f"[DeepSeek 오류: {e}]"

    elif model == "claude":
        if not ant:
            return "[Claude API 키 없음]"
        try:
            r = ant.messages.create(
                model="claude-opus-4-6", max_tokens=4000,   # 2026-04-16 최신
                system="당신은 최고 수준의 분석가이자 전략가입니다. 모든 응답은 한국어로.",
                messages=[{"role": "user", "content": prompt}]
            )
            return r.content[0].text.strip()
        except Exception as e:
            return f"[Claude 오류: {e}]"

    elif model == "grok":
        if not grok_ai:
            return "[Grok API 키 없음]"
        try:
            r = grok_ai.chat.completions.create(
                model="grok-3",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=2000
            )
            return r.choices[0].message.content.strip()
        except Exception as e:
            return f"[Grok 오류: {e}]"

    elif model == "gemini":
        if not gemini_ai:
            return "[Gemini API 키 없음]"
        try:
            r = gemini_ai.chat.completions.create(
                model="gemini-2.5-flash",   # thinking 내장, 2026 최신 stable
                messages=[{"role": "user", "content": prompt}],
                max_tokens=4000
            )
            return r.choices[0].message.content.strip()
        except Exception as e:
            return f"[Gemini 오류: {e}]"

    elif model in ("gpt-4o", "gpt-4.1"):
        if not oai:
            return "[OpenAI API 키 없음]"
        try:
            r = oai.chat.completions.create(
                model="gpt-4.1",   # 2026-04 최신, 1M context, 더 나은 instruction following
                messages=[{"role": "user", "content": prompt}],
                max_tokens=4000,
                temperature=0.4
            )
            return r.choices[0].message.content.strip()
        except Exception as e:
            return f"[GPT-4.1 오류: {e}]"

    elif model == "o3":
        # OpenAI 최상위 추론 모델 — 수학·논리·복잡한 단계적 분석
        if not oai:
            return "[OpenAI API 키 없음]"
        try:
            r = oai.chat.completions.create(
                model="o3",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=5000
                # o-series는 temperature 미지원
            )
            return r.choices[0].message.content.strip()
        except Exception as e:
            return f"[o3 오류: {e}]"

    elif model == "o4-mini":
        # 빠르고 저렴한 추론 — 간단 분석·QA에 최적
        if not oai:
            return "[OpenAI API 키 없음]"
        try:
            r = oai.chat.completions.create(
                model="o4-mini",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=3000
            )
            return r.choices[0].message.content.strip()
        except Exception as e:
            return f"[o4-mini 오류: {e}]"

    return f"[알 수 없는 모델: {model}]"


def _tool_lookup_tool(query: str) -> str:
    """툴 지식 DB에서 검색. 없으면 웹에서 찾아 DB에 추가."""
    tools = _load_tools_kb()
    query_lower = query.lower()

    # 이름 또는 카테고리/용도로 검색
    matches = [
        t for t in tools
        if query_lower in t["name"].lower()
        or query_lower in t.get("category", "").lower()
        or any(query_lower in u.lower() for u in t.get("use_cases", []))
        or query_lower in t.get("description", "").lower()
    ]

    if matches:
        lines = []
        for t in matches[:3]:
            lines.append(f"## {t['name']} ({t['category']})")
            lines.append(t["description"])
            lines.append(f"용도: {', '.join(t.get('use_cases', []))}")
            p = t.get("pricing", {})
            if p.get("free"):  lines.append(f"무료: {p['free']}")
            if p.get("paid"):  lines.append(f"유료: {p['paid']}")
            lines.append(f"강점: {', '.join(t.get('strengths', []))}")
            lines.append(f"주의: {', '.join(t.get('weaknesses', []))}")
            lines.append(f"시작 방법: {t.get('how_to_start', '')}")
            lines.append(f"추천 상황: {t.get('best_for', '')}")
            lines.append("")
        return "\n".join(lines)

    # DB에 없으면 웹 검색으로 보완
    search_result = _tool_web_search(f"{query} 툴 기능 가격 사용법 무료 유료", num_results=3)
    return f"[DB에 없음] 웹 검색 결과:\n{search_result[:1500]}"


def _tool_vibe_coding_guide(idea: str, level: str = "입문") -> str:
    """
    바이브코딩 로드맵 생성.
    아이디어를 받아서 어떤 툴로 어떻게 단계별로 만들지 안내.
    level: 입문 | 중급 | 고급
    """
    tools = _load_tools_kb()
    tools_summary = "\n".join(
        f"- {t['name']} ({t['category']}): {t['description'][:60]}"
        for t in tools
    )

    prompt = f"""당신은 바이브코딩 전문가 멘토입니다.

사용자 아이디어: {idea}
사용자 수준: {level}

사용 가능한 툴 목록:
{tools_summary}

아래 형식으로 실행 가능한 로드맵을 한국어로 작성해주세요:

1. 이 아이디어를 한 줄로 재정의 (더 명확하게)
2. 추천 툴 조합 (왜 이 툴인지 이유 포함)
3. 단계별 구현 순서 (각 단계 예상 시간 포함)
4. 첫 번째로 해야 할 것 딱 하나
5. 막힐 수 있는 포인트와 해결 방법

{level}자도 바로 시작할 수 있게 구체적으로 작성해주세요."""

    return _tool_ask_specialist("claude", prompt)


def _tool_generate_image(prompt: str, size: str = "1024x1024", quality: str = "standard") -> str:
    """DALL-E 3로 이미지 생성 후 저장"""
    if not oai:
        return "[OpenAI API 키 없음]"
    try:
        resp = oai.images.generate(
            model="dall-e-3",
            prompt=prompt,
            size=size,
            quality=quality,
            n=1
        )
        url = resp.data[0].url
        revised = resp.data[0].revised_prompt or prompt

        # 다운로드 & 저장
        import urllib.request as _ur
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        fname = _SAVE_DIR / f"image_{ts}.png"
        _ur.urlretrieve(url, fname)
        return f"이미지 생성 완료\n저장 경로: {fname}\n실제 프롬프트: {revised}"
    except Exception as e:
        return f"[이미지 생성 오류: {e}]"


def _tool_create_chart(chart_type: str, data: dict, title: str = "",
                       x_label: str = "", y_label: str = "",
                       filename: str = "") -> str:
    """
    matplotlib으로 차트/그래프 생성 후 저장.
    chart_type: bar, line, pie, scatter, area
    data: {"labels": [...], "values": [...]} 또는 {"series": [{"name":..,"values":[..]},...], "labels":[...]}
    """
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        import matplotlib.font_manager as fm

        # 한글 폰트 설정
        font_candidates = ["Malgun Gothic", "NanumGothic", "AppleGothic", "sans-serif"]
        for font in font_candidates:
            try:
                plt.rcParams["font.family"] = font
                plt.rcParams["axes.unicode_minus"] = False
                break
            except Exception:
                continue

        fig, ax = plt.subplots(figsize=(10, 6))

        labels = data.get("labels", [])
        values = data.get("values", [])
        series = data.get("series", [])

        ctype = chart_type.lower()

        if ctype == "pie":
            ax.pie(values, labels=labels, autopct="%1.1f%%", startangle=90)
            ax.axis("equal")

        elif ctype == "bar":
            if series:
                x = range(len(labels))
                w = 0.8 / len(series)
                for i, s in enumerate(series):
                    offset = [xi + i * w for xi in x]
                    ax.bar(offset, s["values"], width=w, label=s["name"])
                ax.set_xticks([xi + w * (len(series)-1)/2 for xi in x])
                ax.set_xticklabels(labels)
                ax.legend()
            else:
                ax.bar(labels, values)

        elif ctype == "line":
            if series:
                for s in series:
                    ax.plot(labels, s["values"], marker="o", label=s["name"])
                ax.legend()
            else:
                ax.plot(labels, values, marker="o")

        elif ctype == "area":
            if series:
                for s in series:
                    ax.fill_between(range(len(labels)), s["values"], alpha=0.4, label=s["name"])
                    ax.plot(range(len(labels)), s["values"])
                ax.set_xticks(range(len(labels)))
                ax.set_xticklabels(labels)
                ax.legend()
            else:
                ax.fill_between(range(len(labels)), values, alpha=0.4)
                ax.plot(range(len(labels)), values)
                ax.set_xticks(range(len(labels)))
                ax.set_xticklabels(labels)

        elif ctype == "scatter":
            x_vals = data.get("x", values)
            y_vals = data.get("y", values)
            ax.scatter(x_vals, y_vals)

        if title:   ax.set_title(title, fontsize=14, fontweight="bold")
        if x_label: ax.set_xlabel(x_label)
        if y_label: ax.set_ylabel(y_label)
        ax.grid(True, alpha=0.3)
        plt.tight_layout()

        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        fname_str = filename or f"chart_{ts}.png"
        if not fname_str.endswith(".png"):
            fname_str += ".png"
        save_path = _SAVE_DIR / fname_str
        plt.savefig(save_path, dpi=150, bbox_inches="tight")
        plt.close()
        return f"차트 생성 완료\n저장 경로: {save_path}"
    except Exception as e:
        return f"[차트 생성 오류: {e}]"


def _tool_save_research(formats: list) -> str:
    """
    마지막 리서치 결과를 원하는 형식으로 저장.
    formats: ["excel"], ["pdf"], ["ppt"], ["excel","pdf","ppt"] 등 조합 가능
    """
    if not _last_research:
        return "[저장할 리서치 결과가 없습니다. 먼저 리서치를 실행하세요.]"
    if not _RL:
        return "[리서치 모듈 로드 실패]"

    import re as _re
    from datetime import datetime as _dt

    topic    = _last_research.get("topic", "리서치")
    plan     = _last_research.get("plan", {})
    analysis = _last_research.get("analysis", {})
    results  = _last_research.get("results", [])

    ts   = _dt.now().strftime("%Y%m%d_%H%M")
    safe = _re.sub(r'[^\w가-힣]', '_', topic)[:30]
    base = f"리서치_{safe}_{ts}"
    saved = []

    fmt = [f.lower() for f in formats]

    if "excel" in fmt:
        try:
            p = _rl._save_excel(base, topic, analysis, results)
            if p: saved.append(f"Excel: {p}")
        except Exception as e:
            saved.append(f"Excel 실패: {e}")

    if "pdf" in fmt:
        try:
            p = _rl._save_pdf(base, topic, analysis)
            if p: saved.append(f"PDF: {p}")
        except Exception as e:
            saved.append(f"PDF 실패: {e}")

    if "ppt" in fmt:
        try:
            p = _rl._save_ppt(base, topic, analysis)
            if p: saved.append(f"PPT: {p}")
        except Exception as e:
            saved.append(f"PPT 실패: {e}")

    if not saved:
        return "[저장된 파일 없음 — excel, pdf, ppt 중 하나 이상 지정하세요]"
    return "저장 완료:\n" + "\n".join(saved)


def _tool_web_research(topic: str, depth: str = "중간") -> str:
    """
    기존 웹 리서치 파이프라인 전체 실행.
    make_plan → 웹 수집 → 구조화 추출 → 시장 분석 → 전략 인사이트
    결과를 텍스트 요약으로 반환.
    """
    if not _RL:
        return "[웹 리서치 모듈 로드 실패]"
    try:
        console.print(f"  [dim]  플랜 수립 중...[/dim]")
        plan = _rl.make_plan(topic, depth)
        plan["_depth"] = depth

        console.print(f"  [dim]  웹 수집 중 (쿼리 {len(plan.get('queries', []))}개)...[/dim]")
        results = _rl.run_research(plan)

        console.print(f"  [dim]  {len(results)}개 페이지 분석 중...[/dim]")
        analysis = _rl.analyze(topic, plan, results)

        # 텍스트 요약 구성
        lines = []
        lines.append(f"## 리서치 결과: {topic}")
        lines.append(f"수집 페이지: {len(results)}개\n")

        summary = analysis.get("summary", "")
        if summary:
            lines.append(f"### 시장 현황\n{summary}\n")

        strategy = analysis.get("strategy", "")
        if strategy:
            lines.append(f"### 전략 인사이트\n{strategy}\n")

        per_url = analysis.get("per_url", [])
        if per_url:
            lines.append("### 주요 업체")
            for item in per_url[:8]:
                name = item.get("업체명") or item.get("name") or ""
                url  = item.get("url", "")
                if name:
                    lines.append(f"- {name}  {url}")

        # 결과 캐시 (save_research 도구에서 재사용)
        _last_research.update({
            "topic": topic, "plan": plan,
            "analysis": analysis, "results": results
        })

        lines.append("\n어떤 형식으로 저장할까요? Excel, PDF, PPT 중 원하는 대로 말씀해주세요.")
        return "\n".join(lines)
    except Exception as e:
        return f"[웹 리서치 오류: {e}]"


def _tool_vercel_deploy(project_dir: str, project_name: str = "",
                        framework: str = "") -> str:
    """
    로컬 폴더를 Vercel에 배포. 라이브 URL 반환.
    project_dir : 배포할 폴더 경로 (HTML/CSS/JS 또는 Next.js 등)
    project_name: Vercel 프로젝트 이름 (영문 소문자, 기본: 폴더명)
    framework   : 프레임워크 힌트 (nextjs / react / vue / svelte / 비워두면 정적)
    """
    token = VERCEL_TOKEN or os.getenv("VERCEL_TOKEN", "")
    if not token:
        return (
            "[Vercel 토큰 없음]\n"
            "1. https://vercel.com/account/tokens 에서 토큰 발급\n"
            "2. _local_keys.py 의 VERCEL_TOKEN 에 입력 후 재시작"
        )

    p = Path(project_dir).expanduser().resolve()
    if not p.exists():
        return f"[폴더 없음: {project_dir}]"

    name = (project_name or p.name).lower()
    name = re.sub(r"[^a-z0-9\-]", "-", name)[:50].strip("-") or "maestro-deploy"

    # ── 파일 수집 ─────────────────────────────────────────────────
    files_payload = []
    skip_dirs  = {".git", "node_modules", ".next", "__pycache__", "dist", ".vercel"}
    skip_exts  = {".pyc", ".pyo", ".DS_Store"}
    max_bytes  = 5 * 1024 * 1024   # 5MB 초과 파일 건너뜀

    for fpath in p.rglob("*"):
        if fpath.is_dir():
            continue
        if any(part in skip_dirs for part in fpath.parts):
            continue
        if fpath.suffix in skip_exts:
            continue
        rel = fpath.relative_to(p).as_posix()
        try:
            raw = fpath.read_bytes()
            if len(raw) > max_bytes:
                continue
            import base64 as _b64
            files_payload.append({
                "file": rel,
                "data": _b64.b64encode(raw).decode(),
                "encoding": "base64"
            })
        except Exception:
            continue

    if not files_payload:
        return "[배포할 파일이 없습니다]"

    console.print(f"  [dim]  {len(files_payload)}개 파일 업로드 중...[/dim]")

    # ── 배포 요청 ─────────────────────────────────────────────────
    fw_map = {
        "nextjs": "nextjs", "next": "nextjs",
        "react":  "create-react-app",
        "vue":    "vue",
        "svelte": "svelte",
        "vite":   "vite",
    }
    fw = fw_map.get((framework or "").lower(), None)

    payload: dict = {
        "name":   name,
        "files":  files_payload,
        "target": "production",
        "projectSettings": {}
    }
    if fw:
        payload["projectSettings"]["framework"] = fw

    body = json.dumps(payload, ensure_ascii=False).encode()
    req  = urllib.request.Request(
        "https://api.vercel.com/v13/deployments",
        data=body, method="POST",
        headers={
            "Authorization": f"Bearer {token}",
            "Content-Type":  "application/json",
        }
    )
    try:
        with urllib.request.urlopen(req, timeout=120) as r:
            data = json.loads(r.read())
    except urllib.error.HTTPError as e:
        err = e.read().decode(errors="replace")
        return f"[Vercel API 오류 {e.code}] {err[:300]}"
    except Exception as e:
        return f"[배포 오류: {e}]"

    deploy_id = data.get("id", "")
    url = data.get("url", "")
    alias_url = f"https://{url}" if url and not url.startswith("http") else url

    if not deploy_id:
        return f"[배포 실패] {json.dumps(data, ensure_ascii=False)[:300]}"

    # ── 배포 완료 대기 (최대 3분) ─────────────────────────────────
    console.print(f"  [dim]  배포 중... (최대 3분)[/dim]")
    for _ in range(36):
        time.sleep(5)
        try:
            chk = urllib.request.Request(
                f"https://api.vercel.com/v13/deployments/{deploy_id}",
                headers={"Authorization": f"Bearer {token}"}
            )
            with urllib.request.urlopen(chk, timeout=10) as r2:
                status_data = json.loads(r2.read())
            state = status_data.get("readyState", status_data.get("status", ""))
            if state in ("READY", "ready"):
                live_url = status_data.get("url", url)
                if live_url and not live_url.startswith("http"):
                    live_url = f"https://{live_url}"
                return (
                    f"배포 완료!\n"
                    f"URL: {live_url}\n"
                    f"프로젝트: {name}\n"
                    f"파일: {len(files_payload)}개"
                )
            if state in ("ERROR", "CANCELED"):
                err_msg = status_data.get("errorMessage", "")
                return f"[배포 실패] 상태: {state}  {err_msg}"
        except Exception:
            continue

    # 타임아웃 — URL은 있음
    return (
        f"배포 요청 완료 (상태 확인 타임아웃)\n"
        f"URL: {alias_url}\n"
        f"Vercel 대시보드에서 확인하세요."
    )


def _tool_meeting_to_notion(audio_path: str = "", transcript: str = "",
                            meeting_date: str = "", speaker_map: str = "") -> str:
    """
    회의 음성 파일 또는 텍스트를 분석해 노션에 자동 기록.
    audio_path: 음성 파일 경로 (.mp3/.mp4/.wav/.m4a 등)
    transcript: 직접 텍스트 입력 시 (audio_path 없을 때)
    meeting_date: 회의 날짜 (기본: 오늘)
    speaker_map: 화자 매핑 JSON 문자열 '{"SPEAKER_00":"조경일","SPEAKER_01":"소지민"}'
    """
    if not meeting_date:
        meeting_date = datetime.now().strftime("%Y-%m-%d")

    # 화자 맵 파싱
    spk_map: dict = {}
    if speaker_map:
        try:
            spk_map = json.loads(speaker_map)
        except Exception:
            pass

    # ── 서비스 패키지 import ──────────────────────────────────────
    try:
        from services.transcribe    import transcribe_audio
        from services.diarize       import diarize_transcript
        from services.analyze       import analyze_meeting
        from services.notion_client import push_to_notion
    except ImportError as e:
        return f"[회의록 모듈 로드 실패: {e}]"

    labeled_text = ""

    # ── 음성 파일 처리 ────────────────────────────────────────────
    if audio_path:
        p = Path(audio_path).expanduser()
        if not p.exists():
            return f"[파일 없음: {audio_path}]"

        console.print("  [dim]  STT 변환 중 (Whisper)...[/dim]")
        raw_text = transcribe_audio(str(p))
        if raw_text.startswith("[오류]"):
            return raw_text

        console.print("  [dim]  화자 분리 중 (GPT-4.1-mini)...[/dim]")
        labeled_text = diarize_transcript(raw_text, spk_map if spk_map else None)

    # ── 텍스트 직접 입력 ──────────────────────────────────────────
    elif transcript:
        labeled_text = diarize_transcript(transcript, spk_map if spk_map else None)

    else:
        return "[audio_path 또는 transcript 중 하나는 필수입니다]"

    if not labeled_text.strip():
        return "[변환된 내용이 없습니다]"

    console.print(f"  [dim]  회의 분석 중 (GPT-4.1)...[/dim]")
    analysis = analyze_meeting(labeled_text, meeting_date)

    console.print("  [dim]  노션 페이지 생성 중...[/dim]")
    result = push_to_notion(analysis, labeled_text)

    # 오류가 아니면 성공 요약 추가
    if not result.startswith("[오류]") and not result.startswith("[Notion API 오류"):
        summary = analysis.get("summary", "")
        actions = analysis.get("action_items", [])
        action_str = "\n".join(
            f"  • {a.get('task','')} — {a.get('owner','')} ({a.get('due','')})"
            for a in actions[:5]
        )
        return (
            f"{result}\n\n"
            f"📋 회의 요약: {summary[:200]}\n"
            + (f"\n🎯 액션 아이템:\n{action_str}" if action_str else "")
        )

    return result


def _tool_analyze_document(path: str, question: str = "") -> str:
    """
    파일을 읽고 AI로 분석.
    지원: PDF, Word(.docx), Excel(.xlsx), CSV, 이미지(.png/.jpg/.webp), 텍스트/코드
    question이 있으면 해당 질문에 답하고, 없으면 핵심 내용 요약.
    """
    p = Path(path).expanduser()
    if not p.exists():
        return f"[파일 없음: {path}]"

    ext = p.suffix.lower().lstrip(".")
    text = ""

    # ── PDF ──────────────────────────────────────────────────────
    if ext == "pdf":
        try:
            import pdfplumber
            with pdfplumber.open(str(p)) as pdf:
                pages = [pg.extract_text() or "" for pg in pdf.pages[:40]]
            text = "\n".join(pages)
        except ImportError:
            try:
                from pypdf import PdfReader
                reader = PdfReader(str(p))
                text = "\n".join(pg.extract_text() or "" for pg in reader.pages[:40])
            except ImportError:
                return "[PDF 읽기 실패] pip install pdfplumber 또는 pip install pypdf"
        except Exception as e:
            return f"[PDF 오류: {e}]"

    # ── Word ─────────────────────────────────────────────────────
    elif ext in ("docx", "doc"):
        try:
            from docx import Document
            doc = Document(str(p))
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            # 표도 추출
            for table in doc.tables:
                for row in table.rows:
                    parts.append("\t".join(cell.text for cell in row.cells))
            text = "\n".join(parts)
        except ImportError:
            return "[Word 읽기 실패] pip install python-docx"
        except Exception as e:
            return f"[Word 오류: {e}]"

    # ── Excel ────────────────────────────────────────────────────
    elif ext in ("xlsx", "xls"):
        try:
            import openpyxl
            wb = openpyxl.load_workbook(str(p), read_only=True, data_only=True)
            rows_text = []
            for shname in wb.sheetnames[:5]:
                ws = wb[shname]
                rows_text.append(f"[시트: {shname}]")
                for row in list(ws.iter_rows(values_only=True))[:200]:
                    rows_text.append("\t".join("" if c is None else str(c) for c in row))
            text = "\n".join(rows_text)
        except Exception as e:
            return f"[Excel 오류: {e}]"

    # ── CSV ──────────────────────────────────────────────────────
    elif ext == "csv":
        try:
            import csv as _csv
            lines = []
            for enc in ("utf-8", "cp949", "euc-kr"):
                try:
                    with open(p, encoding=enc, errors="replace") as f:
                        for _, row in zip(range(300), _csv.reader(f)):
                            lines.append("\t".join(row))
                    break
                except Exception:
                    continue
            text = "\n".join(lines)
        except Exception as e:
            return f"[CSV 오류: {e}]"

    # ── 이미지 — GPT-4o Vision ────────────────────────────────────
    elif ext in ("png", "jpg", "jpeg", "webp", "gif", "bmp"):
        if not oai:
            return "[OpenAI API 키 없음]"
        try:
            import base64 as _b64
            with open(p, "rb") as f:
                img_b64 = _b64.b64encode(f.read()).decode()
            mime = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
                    "webp": "image/webp", "gif": "image/gif", "bmp": "image/bmp"}.get(ext, "image/png")
            q = (question or
                 "이 이미지를 상세히 분석해주세요. "
                 "텍스트가 있으면 전부 추출하고, 내용의 의미와 시사점을 설명해주세요.")
            r = oai.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": [
                    {"type": "text", "text": q},
                    {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{img_b64}"}}
                ]}],
                max_tokens=2000
            )
            return r.choices[0].message.content.strip()
        except Exception as e:
            return f"[이미지 분석 오류: {e}]"

    # ── 텍스트 / 코드 ─────────────────────────────────────────────
    else:
        for enc in ("utf-8", "cp949", "euc-kr"):
            try:
                text = p.read_text(encoding=enc, errors="replace")
                break
            except Exception:
                continue
        if not text:
            return f"[읽기 실패: {path}]"

    if not text or not text.strip():
        return "[파일 내용이 비어있거나 텍스트를 추출할 수 없습니다]"

    if not oai:
        return f"[내용 추출 완료 — API 키 없어 분석 불가]\n{text[:3000]}"

    q = (question or
         "이 파일의 핵심 내용을 분석해주세요. "
         "중요한 수치, 인사이트, 시사점을 포함해서 요약해주세요.")
    try:
        r = oai.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content":
                f"파일명: {p.name}\n질문: {q}\n\n내용:\n{text[:9000]}"}],
            max_tokens=2500,
            temperature=0.3
        )
        return r.choices[0].message.content.strip()
    except Exception as e:
        return f"[분석 오류: {e}]\n\n내용 일부:\n{text[:2000]}"


def _tool_ask_claude_code(prompt: str, cwd: str = ".", timeout: int = 180) -> str:
    """
    실제 Claude Code CLI를 subprocess로 호출.
    코딩, 파일 작업, 복잡한 멀티파일 수정에 최적.
    claude -p "prompt" --output-format json --dangerously-skip-permissions
    """
    import shutil
    claude_bin = shutil.which("claude") or "claude"

    work_dir = str(Path(cwd).expanduser().resolve())

    cmd = [
        claude_bin,
        "-p", prompt,
        "--output-format", "json",
        "--dangerously-skip-permissions",
        "--add-dir", work_dir,
    ]

    # Claude Code는 OAuth 세션으로 인증 — ANTHROPIC_API_KEY 제거해야 충돌 없음
    env = os.environ.copy()
    env.pop("ANTHROPIC_API_KEY", None)

    try:
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=timeout,
            cwd=work_dir,
            env=env,
        )
        raw = result.stdout.strip()
        if not raw:
            err = result.stderr.strip()
            return f"[Claude Code 출력 없음] {err[:300]}"

        # JSON 파싱 시도
        try:
            data = json.loads(raw)
            # stream-json 마지막 라인일 수도 있음
            if isinstance(data, list):
                data = data[-1]
            if data.get("is_error"):
                return f"[Claude Code 오류] {data.get('result', raw)}"
            return data.get("result", raw)
        except json.JSONDecodeError:
            # 텍스트 형식 그대로 반환
            return raw

    except subprocess.TimeoutExpired:
        return f"[Claude Code 타임아웃] {timeout}초 초과"
    except FileNotFoundError:
        return "[Claude Code 미설치] 'claude' 명령어를 찾을 수 없습니다."
    except Exception as e:
        return f"[Claude Code 오류: {e}]"


def _exec_tool(name: str, args: dict) -> str:
    """도구 디스패처"""
    dispatch = {
        "read_file":        lambda: _tool_read_file(**args),
        "write_file":       lambda: _tool_write_file(**args),
        "edit_file":        lambda: _tool_edit_file(**args),
        "run_bash":         lambda: _tool_run_bash(**args),
        "glob_search":      lambda: _tool_glob(**args),
        "grep_search":      lambda: _tool_grep(**args),
        "list_dir":         lambda: _tool_list_dir(**args),
        "web_search":       lambda: _tool_web_search(**args),
        "web_fetch":        lambda: _tool_web_fetch(**args),
        "ask_brain":        lambda: _call_brain(args.get("situation", "")),
        "ask_specialist":   lambda: _tool_ask_specialist(
                                args.get("model", "gpt-4o"), args.get("prompt", ""),
                                args.get("category", "")),
        "ask_ensemble":     lambda: _tool_ask_ensemble(
                                args.get("prompt", ""),
                                args.get("category", "strategy_insight")),
        "ask_claude_code":  lambda: _tool_ask_claude_code(
                                args.get("prompt", ""),
                                args.get("cwd", "."),
                                args.get("timeout", 180)),
        "web_research":     lambda: _tool_web_research(
                                args.get("topic", ""),
                                args.get("depth", "중간")),
        "save_research":    lambda: _tool_save_research(
                                args.get("formats", ["excel"])),
        "meeting_to_notion":  lambda: _tool_meeting_to_notion(
                                args.get("audio_path", ""),
                                args.get("transcript", ""),
                                args.get("meeting_date", ""),
                                args.get("speaker_map", "")),
        "analyze_document":   lambda: _tool_analyze_document(
                                args.get("path", ""),
                                args.get("question", "")),
        "lookup_tool":        lambda: _tool_lookup_tool(args.get("query", "")),
        "vibe_coding_guide":  lambda: _tool_vibe_coding_guide(
                                  args.get("idea", ""),
                                  args.get("level", "입문")),
        "generate_image":   lambda: _tool_generate_image(
                                args.get("prompt", ""),
                                args.get("size", "1024x1024"),
                                args.get("quality", "standard")),
        "create_chart":     lambda: _tool_create_chart(
                                args.get("chart_type", "bar"),
                                args.get("data", {}),
                                args.get("title", ""),
                                args.get("x_label", ""),
                                args.get("y_label", ""),
                                args.get("filename", "")),
        "vercel_deploy":    lambda: _tool_vercel_deploy(
                                args.get("project_dir", ""),
                                args.get("project_name", ""),
                                args.get("framework", "")),
    }
    fn = dispatch.get(name)
    if fn:
        return fn()
    return f"[알 수 없는 도구: {name}]"


# ═══════════════════════════════════════════════════════════════
#  GPT-4o Tool Definitions (OpenAI function calling 형식)
# ═══════════════════════════════════════════════════════════════

_TOOL_DEFS = [
    {"type": "function", "function": {
        "name": "read_file",
        "description": "파일 내용을 읽습니다.",
        "parameters": {"type": "object", "properties": {
            "path":   {"type": "string"},
            "offset": {"type": "integer", "description": "시작 줄 번호 (선택)"},
            "limit":  {"type": "integer", "description": "읽을 줄 수 (선택)"}
        }, "required": ["path"]}
    }},
    {"type": "function", "function": {
        "name": "write_file",
        "description": "파일을 생성하거나 전체 내용을 씁니다.",
        "parameters": {"type": "object", "properties": {
            "path":    {"type": "string"},
            "content": {"type": "string"}
        }, "required": ["path", "content"]}
    }},
    {"type": "function", "function": {
        "name": "edit_file",
        "description": "파일에서 특정 문자열을 찾아 교체합니다. old_str는 파일에서 정확히 일치해야 합니다.",
        "parameters": {"type": "object", "properties": {
            "path":    {"type": "string"},
            "old_str": {"type": "string"},
            "new_str": {"type": "string"}
        }, "required": ["path", "old_str", "new_str"]}
    }},
    {"type": "function", "function": {
        "name": "run_bash",
        "description": "터미널 명령어를 실행합니다.",
        "parameters": {"type": "object", "properties": {
            "command": {"type": "string"},
            "timeout": {"type": "integer"}
        }, "required": ["command"]}
    }},
    {"type": "function", "function": {
        "name": "glob_search",
        "description": "glob 패턴으로 파일을 검색합니다. 예: **/*.py",
        "parameters": {"type": "object", "properties": {
            "pattern": {"type": "string"},
            "path":    {"type": "string", "description": "검색 시작 경로 (기본: 현재 디렉토리)"}
        }, "required": ["pattern"]}
    }},
    {"type": "function", "function": {
        "name": "grep_search",
        "description": "파일 내용에서 패턴을 검색합니다.",
        "parameters": {"type": "object", "properties": {
            "pattern":      {"type": "string"},
            "path":         {"type": "string"},
            "file_pattern": {"type": "string", "description": "파일 필터 (예: *.py)"}
        }, "required": ["pattern"]}
    }},
    {"type": "function", "function": {
        "name": "list_dir",
        "description": "디렉토리 내용을 나열합니다.",
        "parameters": {"type": "object", "properties": {
            "path": {"type": "string"}
        }, "required": ["path"]}
    }},
    {"type": "function", "function": {
        "name": "web_search",
        "description": "웹에서 정보를 검색합니다.",
        "parameters": {"type": "object", "properties": {
            "query":       {"type": "string"},
            "num_results": {"type": "integer"}
        }, "required": ["query"]}
    }},
    {"type": "function", "function": {
        "name": "web_fetch",
        "description": "URL의 페이지 내용을 가져옵니다.",
        "parameters": {"type": "object", "properties": {
            "url": {"type": "string"}
        }, "required": ["url"]}
    }},
    {"type": "function", "function": {
        "name": "ask_brain",
        "description": (
            "조경일 뇌 에이전트에게 마케팅/전략/비즈니스 판단을 요청합니다. "
            "전략적 결정, 마케팅 방향, 우선순위 판단 등에 사용하세요."
        ),
        "parameters": {"type": "object", "properties": {
            "situation": {"type": "string", "description": "판단이 필요한 상황 설명"}
        }, "required": ["situation"]}
    }},
    {"type": "function", "function": {
        "name": "ask_specialist",
        "description": (
            "특정 LLM 전문가에게 작업을 위임합니다. 2026년 4월 최신 모델 기준.\n"
            "- o3      : 최상위 추론. 수학·논리·복잡한 단계별 분석 (가장 강력, 느림)\n"
            "- claude  : claude-opus-4-6. 긴 문서·보고서·전략 인사이트·세밀한 글쓰기\n"
            "- gemini  : gemini-2.5-flash (thinking 내장). 멀티모달·대용량 컨텍스트·문서 요약\n"
            "- gpt-4.1 : 구조화 출력·데이터 추출·코딩·instruction 정밀 수행\n"
            "- deepseek: 복잡한 추론·알고리즘·코드 분석 (저비용 추론)\n"
            "- grok    : 실시간/최신 정보, 뉴스, 트렌드, X(Twitter) 데이터\n"
            "- o4-mini : 빠른 추론·간단 분석·QA (저비용)\n"
            "- auto    : 오케스트레이터가 category 기준으로 최적 모델 자동 선택"
        ),
        "parameters": {"type": "object", "properties": {
            "model":    {"type": "string",
                         "enum": ["o3", "claude", "gemini", "gpt-4.1", "deepseek",
                                  "grok", "o4-mini", "auto"]},
            "prompt":   {"type": "string"},
            "category": {"type": "string",
                         "description": (
                             "model=auto일 때 필수 — 오케스트레이터 DB가 이 카테고리 기준으로 현재 최강 모델을 동적 선택함. "
                             "하드코딩 금지: 오늘의 최강 모델이 내일 바뀔 수 있음."
                         ),
                         "enum": [
                             "query_generation", "url_filtering", "data_extraction",
                             "market_analysis", "strategy_insight",
                             "document_writing", "document_export",
                             "coding", "quick_qa", "deep_reasoning", "realtime_info"
                         ]}
        }, "required": ["model", "prompt"]}
    }},
    {"type": "function", "function": {
        "name": "ask_ensemble",
        "description": (
            "Claude + Gemini + DeepSeek 를 병렬 호출한 뒤 GPT-4.1 이 결과를 종합합니다. "
            "현존 최고 품질이 필요할 때 사용. 단일 모델의 blind spot을 상호 보완. "
            "전략 분석·시장 조사·중요 보고서 작성·복잡한 의사결정에 적합."
        ),
        "parameters": {"type": "object", "properties": {
            "prompt":   {"type": "string", "description": "앙상블에게 위임할 상세 프롬프트"},
            "category": {"type": "string",
                         "description": "작업 카테고리 (오케스트레이터 DB 학습에 사용)",
                         "enum": [
                             "query_generation", "url_filtering", "data_extraction",
                             "market_analysis", "strategy_insight",
                             "document_writing", "document_export",
                             "coding", "quick_qa", "deep_reasoning", "realtime_info"
                         ]}
        }, "required": ["prompt"]}
    }},
    {"type": "function", "function": {
        "name": "meeting_to_notion",
        "description": (
            "회의 음성 파일 또는 회의 텍스트를 분석해 노션에 자동으로 회의록을 작성합니다. "
            "음성 파일(.mp3/.mp4/.wav/.m4a)을 주면 STT→화자분리→분석→노션 전체 자동 처리. "
            "'회의록 정리해줘', '회의 내용 노션에 올려줘' 같은 요청에 사용하세요."
        ),
        "parameters": {"type": "object", "properties": {
            "audio_path":   {"type": "string", "description": "음성 파일 전체 경로 (mp3/mp4/wav/m4a)"},
            "transcript":   {"type": "string", "description": "텍스트 직접 입력 (audio_path 없을 때). 형식: '화자: 내용\\n화자: 내용'"},
            "meeting_date": {"type": "string", "description": "회의 날짜 (예: 2026-04-17). 기본: 오늘"},
            "speaker_map":  {"type": "string", "description": "화자 이름 매핑 JSON '{\"SPEAKER_00\":\"조경일\",\"SPEAKER_01\":\"소지민\"}'"}
        }, "required": []}
    }},
    {"type": "function", "function": {
        "name": "analyze_document",
        "description": (
            "파일을 읽고 AI로 분석합니다. "
            "PDF, Word(.docx), Excel(.xlsx), CSV, 이미지(.png/.jpg/.webp), 텍스트/코드 파일 모두 지원. "
            "사용자가 파일 경로를 주거나 '이 파일 분석해줘', '문서 요약해줘'라고 하면 사용하세요. "
            "이미지는 GPT-4o Vision으로 텍스트 추출 및 내용 분석까지 가능합니다."
        ),
        "parameters": {"type": "object", "properties": {
            "path":     {"type": "string", "description": "분석할 파일의 전체 경로"},
            "question": {"type": "string", "description": "파일에 대해 묻고 싶은 구체적인 질문 (없으면 전체 요약)"}
        }, "required": ["path"]}
    }},
    {"type": "function", "function": {
        "name": "lookup_tool",
        "description": "툴/서비스 정보를 조회합니다. 어떤 툴인지, 기능, 가격, 무료/유료, 시작 방법을 알 수 있습니다. Carrd, Framer, n8n, Supabase, Vercel, Notion, Readdy, Antigravity, Codx 등.",
        "parameters": {"type": "object", "properties": {
            "query": {"type": "string", "description": "툴 이름 또는 용도 (예: '랜딩페이지', 'Carrd', '자동화')"}
        }, "required": ["query"]}
    }},
    {"type": "function", "function": {
        "name": "vibe_coding_guide",
        "description": "바이브코딩 로드맵을 만들어줍니다. 아이디어를 말하면 어떤 툴로 어떻게 단계별로 만들지 구체적인 계획을 제시합니다. 막막한 사람, 처음 시작하는 사람에게 특히 유용합니다.",
        "parameters": {"type": "object", "properties": {
            "idea":  {"type": "string", "description": "만들고 싶은 것"},
            "level": {"type": "string", "enum": ["입문", "중급", "고급"], "description": "사용자 기술 수준"}
        }, "required": ["idea"]}
    }},
    {"type": "function", "function": {
        "name": "generate_image",
        "description": "DALL-E 3로 이미지를 생성합니다. 마케팅 배너, 썸네일, 일러스트, 아이디어 시각화 등에 사용하세요.",
        "parameters": {"type": "object", "properties": {
            "prompt":  {"type": "string", "description": "이미지 설명 (구체적일수록 좋음)"},
            "size":    {"type": "string", "enum": ["1024x1024", "1792x1024", "1024x1792"],
                        "description": "정사각형|가로형|세로형"},
            "quality": {"type": "string", "enum": ["standard", "hd"]}
        }, "required": ["prompt"]}
    }},
    {"type": "function", "function": {
        "name": "create_chart",
        "description": "데이터를 차트/그래프로 시각화합니다. bar(막대), line(선), pie(파이), area(면적), scatter(산점도) 지원.",
        "parameters": {"type": "object", "properties": {
            "chart_type": {"type": "string", "enum": ["bar", "line", "pie", "area", "scatter"]},
            "data": {
                "type": "object",
                "description": "단일 시리즈: {labels:[...], values:[...]} / 복수 시리즈: {labels:[...], series:[{name:..,values:[..]},..]}",
            },
            "title":    {"type": "string"},
            "x_label":  {"type": "string"},
            "y_label":  {"type": "string"},
            "filename": {"type": "string", "description": "저장 파일명 (확장자 없이)"}
        }, "required": ["chart_type", "data"]}
    }},
    {"type": "function", "function": {
        "name": "save_research",
        "description": (
            "리서치 결과를 원하는 형식으로 저장합니다. "
            "web_research 실행 후 사용자가 저장 형식을 말하면 호출하세요. "
            "Excel, PDF, PPT 중 원하는 것을 복수로 지정할 수 있습니다."
        ),
        "parameters": {"type": "object", "properties": {
            "formats": {
                "type": "array",
                "items": {"type": "string", "enum": ["excel", "pdf", "ppt"]},
                "description": "저장할 형식 목록. 예: ['excel'], ['excel','pdf'], ['excel','pdf','ppt']"
            }
        }, "required": ["formats"]}
    }},
    {"type": "function", "function": {
        "name": "web_research",
        "description": (
            "웹 리서치 전체 파이프라인을 실행합니다. "
            "경쟁사 분석, 시장 조사, 업체 비교, 트렌드 파악 등 "
            "특정 주제에 대해 웹을 자동 수집하고 구조화 분석까지 완료합니다. "
            "결과는 자동으로 파일로 저장됩니다."
        ),
        "parameters": {"type": "object", "properties": {
            "topic": {"type": "string", "description": "조사할 주제"},
            "depth": {"type": "string",
                      "description": "조사 깊이: '빠른' (5개 URL), '중간' (10개), '깊은' (20개)",
                      "enum": ["빠른", "중간", "깊은"]}
        }, "required": ["topic"]}
    }},
    {"type": "function", "function": {
        "name": "ask_claude_code",
        "description": (
            "【현존 최강 코딩 에이전트】 Claude Code CLI를 호출합니다. "
            "코드 작성·버그 수정·리팩터링·멀티파일 구현·테스트 실행·앱 전체 개발 → 반드시 이 도구 먼저. "
            "ask_specialist(claude)로 코딩하지 마세요 — Claude Code가 파일 접근+실행까지 통합되어 10배 강합니다. "
            "사용 예: '앱 만들어줘', '이 버그 고쳐줘', '함수 짜줘', 'API 서버 구현', '리팩터링해줘'."
        ),
        "parameters": {"type": "object", "properties": {
            "prompt":  {"type": "string", "description": "Claude Code에게 전달할 작업 지시"},
            "cwd":     {"type": "string", "description": "작업 디렉토리 경로 (기본: 현재 폴더)"},
            "timeout": {"type": "integer", "description": "최대 대기 시간(초, 기본 180)"}
        }, "required": ["prompt"]}
    }},
    {"type": "function", "function": {
        "name": "vercel_deploy",
        "description": (
            "로컬 폴더를 Vercel에 배포하고 라이브 URL을 반환합니다. "
            "HTML/CSS/JS 정적 사이트, Next.js, React, Vue, Svelte 등 지원. "
            "랜딩페이지, 포트폴리오, 웹앱 등을 바로 배포할 때 사용하세요. "
            "VERCEL_TOKEN이 없으면 발급 안내를 합니다."
        ),
        "parameters": {"type": "object", "properties": {
            "project_dir":  {"type": "string", "description": "배포할 폴더 경로"},
            "project_name": {"type": "string", "description": "Vercel 프로젝트 이름 (영문 소문자, 기본: 폴더명)"},
            "framework":    {"type": "string", "description": "프레임워크: nextjs / react / vue / svelte / vite / 비워두면 정적 HTML"}
        }, "required": ["project_dir"]}
    }},
]


# ═══════════════════════════════════════════════════════════════
#  SYSTEM PROMPT
# ═══════════════════════════════════════════════════════════════

def _build_system_prompt() -> str:
    """집단지성 패턴을 포함한 시스템 프롬프트 동적 생성"""
    knowledge_section = ""
    if _PC:
        try:
            knowledge_section = _pc.build_knowledge_prompt()
        except Exception:
            pass
    return _SYSTEM_BASE + knowledge_section


_SYSTEM_BASE = """[절대 규칙] 모든 응답은 반드시 한국어로만 작성합니다. 영어 금지. 전문 용어도 한국어로 설명합니다.

당신은 MAESTRO입니다.

아이디어를 실제로 만들어주고, 방향을 잡아주고, 함께 배워가는 AI입니다.
단순히 답을 주는 게 아니라 — 실제로 실행하고, 가르치고, 함께 성장합니다.

---

## MAESTRO의 핵심 철학

**"다른 사람들은 이런 상황에서 이렇게 했어요"**
비슷한 상황의 사용자들이 발견한 좋은 방법을 자연스럽게 안내합니다.
정답을 강요하지 않고, 선택지와 맥락을 줍니다.

**"같이 만들면서 배워요"**
바이브코딩은 혼자 공부하는 게 아닙니다. 실제로 만들면서 배우는 겁니다.
막막하면 작게 쪼개서 시작합니다. 첫 번째 한 걸음만 제시합니다.

**"어떤 툴을 써야 할지 알고 있어요"**
Carrd, Framer, n8n, Supabase, Vercel, Notion, Claude Code 등
각 툴의 강점과 과금 방식을 알고 있고, 상황에 맞는 것을 추천합니다.

---

## 도구 목록

**vercel_deploy** — Vercel 무료 배포
로컬 폴더를 Vercel에 배포하고 라이브 URL을 즉시 반환합니다.
HTML/CSS/JS 정적 사이트, Next.js, React, Vue, Svelte 모두 지원.
"이 폴더 배포해줘", "랜딩페이지 올려줘" 같은 요청에 사용하세요.
VERCEL_TOKEN이 없으면 발급 경로를 안내합니다.

> **[자동 제안 규칙]** 대화 흐름에서 아래 상황이 감지되면, 먼저 물어보세요:
> - 랜딩페이지 / 포트폴리오 / 소개페이지 / 웹사이트 제작 완료 시점
> - HTML/CSS/JS 파일을 write_file로 생성한 직후
> - "공유하고 싶어", "링크 보내고 싶어", "온라인에 올리고 싶어" 등 배포 의도 발언
> - 클라이언트에게 결과물을 보여줘야 하는 상황
>
> 이때 이렇게 제안하세요:
> "Vercel로 지금 바로 배포해서 라이브 링크 드릴 수 있어요. 해드릴까요?"

**meeting_to_notion** — 회의록 자동화
음성 파일 경로를 주면 STT→화자 분리→AI 분석→노션 페이지 생성까지 전부 자동.
회의 내용을 텍스트로 붙여줘도 됩니다.

**analyze_document** — 파일/문서/이미지 분석
PDF, Word, Excel, CSV, 이미지, 코드 파일 전부 분석 가능.
"이 파일 요약해줘", "이 이미지 텍스트 추출해줘", "엑셀 데이터 분석해줘" 같은 요청에 사용.
파일 경로를 경로 그대로 받아서 바로 처리합니다.

**vibe_coding_guide** — 바이브코딩 로드맵 생성
아이디어가 있는데 어디서 시작해야 할지 모를 때. 어떤 툴로 어떻게 만들지 단계별 계획을 드립니다.

**lookup_tool** — 툴/서비스 정보 조회
"Carrd가 뭐야?", "n8n 무료야?", "랜딩페이지 만드는 툴 뭐 있어?" 같은 질문에 답합니다.

**ask_claude_code** — 실제 Claude Code로 구현
만들어달라고 하면 실제로 파일을 만들고, 코드를 짜고, 실행까지 합니다.

**web_research** — 시장/경쟁사/정보 수집
특정 주제를 깊게 조사해서 구조화된 결과를 줍니다. 완료 후 저장 형식을 물어보세요.

**save_research** — 리서치 결과 저장
Excel, PDF, PPT 중 원하는 형식으로. "다 해줘"하면 세 가지 모두 저장합니다.

**generate_image** — DALL-E 3 이미지 생성
마케팅 배너, 썸네일, 아이디어 시각화 등.

**create_chart** — 차트/그래프 생성
숫자 데이터가 있으면 시각화를 먼저 제안하세요.

**ask_brain** — 조경일 마케터 뇌 판단
마케팅, 전략, 비즈니스 방향 판단. 전략적 선택이 필요할 때 자연스럽게 활용합니다.

**ask_specialist("deepseek")** — 복잡한 추론/알고리즘
**ask_specialist("claude")** — 정밀한 글쓰기/문서 분석
**ask_specialist("grok")** — 실시간/최신 정보

**web_search / web_fetch / read_file / write_file / edit_file / run_bash / glob_search / grep_search / list_dir**
— 검색, 파일, 터미널 직접 접근.

---

## 바이브코딩 가이드 원칙

처음 시작하는 사람이 오면:
1. 아이디어를 먼저 명확하게 만들어줍니다
2. 어떤 툴 조합이 맞는지 이유와 함께 추천합니다
3. 단계를 잘게 쪼개서 첫 번째 것만 시작하게 합니다
4. 각 단계마다 "왜 이렇게 하는지" 설명합니다
5. 막히면 다른 방법을 제시합니다

중급/고급자에게는 설명을 줄이고 실행에 집중합니다.

---

## 도구 라우팅 규칙 (반드시 준수)

당신(GPT-4.1)은 오케스트레이터입니다. 사용자 요청을 깊게 분석하고 최적의 도구를 선택해 결과를 수집하고 종합해서 답변합니다. 도구 없이 혼자 추론만 하는 것은 금지입니다.

---

### 핵심 원칙: 모델 능력은 하드코딩하지 않는다

LLM의 능력은 매주 바뀝니다. "Claude가 코딩 최강" "GPT가 데이터 추출 최강" 같은 고정 판단을 금지합니다.
**모델 선택은 항상 `model="auto"`** — 오케스트레이터 DB가 실시간 평가 기반으로 현재 최강 모델을 선택합니다.

---

### 작업 유형별 도구 선택

**① 실제 파일 접근 · 코드 실행 · 개발 환경이 필요한 작업**
→ `ask_claude_code` 사용 (이건 모델 능력이 아니라 **도구 능력** — 파일 읽기·쓰기·실행 환경 통합)
"파일 만들어줘", "실행해봐", "코드 테스트해", "앱 배포해" → ask_claude_code

**② 코드 작성·버그 분석·알고리즘 설계 (파일 실행 불필요)**
→ `ask_specialist(model="auto", category="coding")`
오케스트레이터 DB가 현재 시점 코딩 최강 모델을 선택 (오늘은 Claude일 수 있고, 3일 후엔 Gemini일 수 있음)

**③ 전략·시장분석·보고서·제안서 (최고 품질 필요)**
→ `web_research` (데이터 수집) → `ask_ensemble` (현재 최강 3개 모델 병렬 → 종합)
앙상블은 blind spot을 상호 보완해 단일 모델보다 품질이 높습니다.

**④ 시장조사·경쟁사·업체 목록**
→ `web_research` 먼저 → `ask_ensemble(category="market_analysis")` → 저장 형식(Excel/PDF/PPT) 질문

**⑤ 전략/마케팅 판단**
→ `ask_brain` (조경일 뇌 에이전트) → `ask_ensemble(category="strategy_insight")` 순서로 결합

**⑥ 문서 작성 (보고서·제안서·긴 글)**
→ `ask_specialist(model="auto", category="document_writing")`

**⑦ 문서 파일 변환 (Excel·PDF·PPT·Word 저장)**
→ `ask_specialist(model="auto", category="document_export")`

**⑧ 실시간·최신 뉴스·트렌드**
→ `ask_specialist(model="auto", category="realtime_info")` 또는 `web_search`

**⑨ 복잡한 수학·논리·단계별 추론**
→ `ask_specialist(model="auto", category="deep_reasoning")`

**⑩ 빠른 QA·단순 질문**
→ `ask_specialist(model="auto", category="quick_qa")`

**⑪ 대용량 문서·긴 컨텍스트·멀티모달**
→ `ask_specialist(model="auto", category="market_analysis")` (컨텍스트 처리 능력 기준 DB 선택)

---

**도구 실패 시 필수 행동** → 실패 사실을 사용자에게 반드시 명시. "도구가 실패해 학습 데이터로 답했습니다"라고 밝힐 것. 오류를 숨기고 아는 척 금지.

**응답 구체성 원칙**
- 숫자가 있으면 숫자로. "상당히", "많이" 대신 구체적 수치/이름/URL 포함.
- 전문가 분석은 bullet 3개 이상, 각 항목에 근거 포함.
- 리서치 결과는 업체명/가격/특징 테이블 형태로 정리.

## 응답 방식
- 반드시 한국어로만. 영어 절대 금지.
- 전문가 결과가 영어로 와도 번역해서 전달합니다.
- 설명보다 실행 먼저. 완료 후 다음을 묻습니다.
- 숫자/데이터 나오면 차트 시각화를 제안합니다.
- 막막해하는 사람에게는 "일단 이것만 해봐요"로 시작합니다."""


# ═══════════════════════════════════════════════════════════════
#  AGENT LOOP — ReAct 패턴
# ═══════════════════════════════════════════════════════════════

def run_agent(user_input: str, history: list, auto_confirm: bool = False) -> str:
    """
    GPT-4o 기반 ReAct 에이전트 루프.
    도구를 반복 호출하며 작업을 완료한 후 최종 답변 반환.
    """
    if not oai:
        return "[OpenAI API 키가 없어 MAESTRO를 실행할 수 없습니다]"

    messages = [{"role": "system", "content": _build_system_prompt()}]
    messages += history[-40:]   # 로컬엔 전체 보존, GPT-4o엔 최근 20턴(40개 메시지)

    # 영어 입력이어도 한국어 응답 강제
    content = user_input
    if user_input and not any("\uAC00" <= c <= "\uD7A3" for c in user_input):
        content = user_input + "\n(반드시 한국어로 답변)"

    # ── 복잡도 판단: 오케스트레이터는 항상 실행, 첫 도구만 유도 ──
    tier = _judge_query_tier(user_input)
    # complex: 전략/분석/보고서 → 첫 도구를 ask_ensemble 로 강제
    # 코딩은 ask_claude_code 로 강제 (앙상블보다 Claude Code가 우위)
    _force_first_tool: str | None = None
    if tier == "complex":
        t_lower = user_input.lower()
        coding_kw = ["코드", "코딩", "개발", "구현", "버그", "함수", "클래스", "앱", "api", "짜줘", "만들어"]
        if any(k in t_lower for k in coding_kw):
            _force_first_tool = "ask_claude_code"
            console.print("  [dim][TIER:complex/coding] Claude Code 우선 실행[/dim]")
        else:
            _force_first_tool = "ask_ensemble"
            console.print("  [dim][TIER:complex] 앙상블 우선 실행 (Claude+Gemini+DeepSeek→GPT-4.1)[/dim]")

    # ── 응답 캐시: 유사도 97% 이상만, user 메시지 끝에 참고로 첨부 ──
    if not history:   # 첫 턴에만 (문맥 없을 때)
        cached_resp, sim = _cache_lookup(user_input)
        if cached_resp:
            content += (
                f"\n\n[이전 유사 질문(유사도 {sim:.0%}) 참고 — 현재 질문과 다르거나 정보가 업데이트됐으면 무시하세요]\n"
                f"{cached_resp[:600]}"
            )
            console.print(f"  [dim][CACHE] 유사 답변 참고 첨부 (유사도 {sim:.0%})[/dim]")

    messages.append({"role": "user", "content": content})

    max_iterations = 20
    iteration = 0

    while iteration < max_iterations:
        iteration += 1

        # 첫 iteration에서 강제 도구가 지정된 경우 tool_choice로 강제
        if iteration == 1 and _force_first_tool:
            tc = {"type": "function", "function": {"name": _force_first_tool}}
        else:
            tc = "auto"

        try:
            response = oai.chat.completions.create(
                model="gpt-4.1",
                messages=messages,
                tools=_TOOL_DEFS,
                tool_choice=tc,
                max_tokens=4000,
                temperature=0.7
            )
        except Exception as e:
            return f"[MAESTRO 오류: {e}]"

        msg = response.choices[0].message

        # 도구 호출 없음 → 최종 답변
        if not msg.tool_calls:
            final = msg.content or ""
            # 첫 턴 + 도구 미사용 응답만 캐시 저장
            # (웹 검색/파일/리서치 결과는 캐시 안 함 — 항상 최신 데이터 필요)
            if iteration == 1 and not history and final and len(final) > 50:
                import threading as _t
                _t.Thread(target=_cache_store, args=(user_input, final), daemon=True).start()
            return final

        # 도구 호출 처리
        messages.append({"role": "assistant",
                         "content": msg.content,
                         "tool_calls": [tc.model_dump() for tc in msg.tool_calls]})

        for tc in msg.tool_calls:
            tool_name = tc.function.name
            try:
                tool_args = json.loads(tc.function.arguments)
            except Exception:
                tool_args = {}

            # bash 명령 확인 (위험 아닐 때도 표시)
            if tool_name == "run_bash" and not auto_confirm:
                cmd = tool_args.get("command", "")
                console.print(f"\n[yellow]  bash>[/yellow] {cmd}")
                if not Confirm.ask("  실행할까요?", default=True):
                    result = "[사용자가 실행 취소]"
                    messages.append({
                        "role": "tool",
                        "tool_call_id": tc.id,
                        "content": result
                    })
                    continue

            # 도구 실행 & 결과 표시
            _show_tool_call(tool_name, tool_args)
            result = _exec_tool(tool_name, tool_args)
            _show_tool_result(tool_name, result)

            messages.append({
                "role": "tool",
                "tool_call_id": tc.id,
                "content": result
            })

    return "[최대 반복 횟수 초과. 작업이 너무 복잡합니다.]"


# ═══════════════════════════════════════════════════════════════
#  UI 헬퍼
# ═══════════════════════════════════════════════════════════════

_TOOL_ICONS = {
    "read_file":       "[READ]",
    "write_file":      "[WRITE]",
    "edit_file":       "[EDIT]",
    "run_bash":        "[BASH]",
    "glob_search":     "[GLOB]",
    "grep_search":     "[GREP]",
    "list_dir":        "[DIR]",
    "web_search":      "[WEB]",
    "web_fetch":       "[FETCH]",
    "ask_brain":       "[BRAIN]",
    "ask_specialist":  "[LLM]",
    "ask_ensemble":    "[ENSEMBLE]",
    "ask_claude_code": "[CLAUDE CODE]",
    "web_research":    "[RESEARCH]",
    "save_research":   "[SAVE]",
    "generate_image":    "[IMAGE]",
    "create_chart":      "[CHART]",
    "meeting_to_notion": "[MEETING]",
    "analyze_document":  "[DOC]",
    "lookup_tool":       "[TOOL]",
    "vibe_coding_guide": "[VIBE]",
    "vercel_deploy":     "[VERCEL]",
}

def _show_tool_call(name: str, args: dict):
    icon = _TOOL_ICONS.get(name, "🔧")
    if name == "read_file":
        console.print(f"  {icon} [dim]읽기:[/dim] {args.get('path', '')}")
    elif name == "write_file":
        console.print(f"  {icon} [dim]쓰기:[/dim] {args.get('path', '')}")
    elif name == "edit_file":
        console.print(f"  {icon} [dim]수정:[/dim] {args.get('path', '')}")
    elif name == "run_bash":
        console.print(f"  {icon} [dim]실행:[/dim] {args.get('command', '')[:80]}")
    elif name == "web_search":
        console.print(f"  {icon} [dim]검색:[/dim] {args.get('query', '')}")
    elif name == "web_fetch":
        console.print(f"  {icon} [dim]페이지:[/dim] {args.get('url', '')[:60]}")
    elif name == "ask_brain":
        console.print(f"  {icon} [dim]뇌 에이전트 판단 중...[/dim]")
    elif name == "ask_specialist":
        model = args.get("model", "")
        cat   = args.get("category", "")
        label = f"{model}({cat})" if model == "auto" and cat else model
        console.print(f"  {icon} [dim]{label} 전문가에게 위임 중...[/dim]")
    elif name == "ask_ensemble":
        cat = args.get("category", "")
        console.print(f"  {icon} [bold magenta]앙상블(Claude+Gemini+DeepSeek→GPT-4.1){' ['+cat+']' if cat else ''} 시작[/bold magenta]")
    elif name == "ask_claude_code":
        cwd = args.get("cwd", ".")
        preview = args.get("prompt", "")[:60]
        console.print(f"  {icon} [bold cyan]{preview}...[/bold cyan]")
        console.print(f"    [dim]작업 경로: {cwd}[/dim]")
    elif name == "web_research":
        topic = args.get("topic", "")
        depth = args.get("depth", "중간")
        console.print(f"  {icon} [bold green]{topic}[/bold green]  [{depth}]")
    elif name == "save_research":
        fmts = ", ".join(args.get("formats", []))
        console.print(f"  {icon} [dim]{fmts} 저장 중...[/dim]")
    elif name == "meeting_to_notion":
        src = args.get("audio_path") or args.get("transcript", "")[:30]
        dt  = args.get("meeting_date", "오늘")
        console.print(f"  [MEETING] [bold cyan]{dt}[/bold cyan] [dim]{src[:40]}[/dim]")
    elif name == "analyze_document":
        fname = Path(args.get("path", "")).name
        q     = args.get("question", "")
        console.print(f"  {icon} [bold yellow]{fname}[/bold yellow]" +
                      (f"  [dim]{q[:40]}[/dim]" if q else "  [dim]전체 분석[/dim]"))
    elif name == "lookup_tool":
        console.print(f"  {icon} [dim]{args.get('query', '')} 조회 중...[/dim]")
    elif name == "vibe_coding_guide":
        console.print(f"  {icon} [bold yellow]{args.get('idea', '')[:50]}...[/bold yellow]")
    elif name == "generate_image":
        p = args.get("prompt", "")[:50]
        console.print(f"  {icon} [bold magenta]{p}...[/bold magenta]")
    elif name == "create_chart":
        ct = args.get("chart_type", "")
        t  = args.get("title", "차트")
        console.print(f"  {icon} [bold green]{ct} - {t}[/bold green]")
    elif name == "vercel_deploy":
        d = args.get("project_dir", "")
        n = args.get("project_name", "") or Path(d).name
        console.print(f"  {icon} [bold cyan]{n}[/bold cyan] [dim]배포 중...[/dim]")
    elif name in ("glob_search", "grep_search"):
        console.print(f"  {icon} [dim]검색:[/dim] {args.get('pattern', '')}")
    elif name == "list_dir":
        console.print(f"  {icon} [dim]목록:[/dim] {args.get('path', '')}")


def _safe(text: str) -> str:
    """cp949로 표현 불가한 문자 제거"""
    return text.encode("cp949", errors="replace").decode("cp949")

def _show_tool_result(name: str, result: str):
    lines = result.count("\n") + 1
    if "[오류]" in result or "[실패]" in result or "[없음]" in result:
        preview = _safe(result[:100].replace("\n", " "))
        console.print(f"    [red]>> {preview}[/red]")
    else:
        console.print(f"    [dim]>> {lines}줄 반환[/dim]")


# ═══════════════════════════════════════════════════════════════
#  메인 터미널 UI
# ═══════════════════════════════════════════════════════════════

def main():
    # 자동 업데이트 (가장 먼저 실행 — 재시작 후 새 버전이 뜸)
    _check_update()

    console.clear()

    # 연결된 모델 확인
    models = []
    if oai:       models.append("GPT-4.1 / o3 / o4-mini")
    if ant:       models.append("Claude Opus 4.6")
    if deepseek:  models.append("DeepSeek-R2")
    if grok_ai:   models.append("Grok-3")
    if gemini_ai: models.append("Gemini-2.5-Flash")

    # 뇌 에이전트 핑 (콜드 스타트 워밍업 포함)
    brain_ok = False
    try:
        ping_url = _BRAIN_URL.replace("/api/research", "/")
        with urllib.request.urlopen(ping_url, timeout=10) as r:
            brain_ok = r.status < 500
    except Exception:
        brain_ok = False

    # 콜드 스타트 대비: 백그라운드에서 미리 /api/research 워밍업
    if brain_ok:
        import threading
        def _warmup():
            try:
                body = json.dumps({"situation": "워밍업"}, ensure_ascii=False).encode()
                req  = urllib.request.Request(_BRAIN_URL, data=body, method="POST",
                                              headers={"Content-Type": "application/json"})
                urllib.request.urlopen(req, timeout=60)
            except Exception:
                pass
        threading.Thread(target=_warmup, daemon=True).start()

    brain_str = "[green]연동됨[/green]" if brain_ok else "[red]미연동[/red]"
    orch_str  = "[green]활성[/green]"   if _ORCH     else "[red]비활성[/red]"

    console.print(Panel(
        f"[bold cyan]MAESTRO[/bold cyan]  v{VERSION}\n"
        f"[dim]현존 최강 AI 오케스트레이터[/dim]\n\n"
        f"  LLM    : {' / '.join(models) if models else '[red]없음[/red]'}\n"
        f"  뇌 에이전트: {brain_str}\n"
        f"  오케스트레이터: {orch_str}",
        border_style="bright_magenta", padding=(1, 6)
    ))

    if not oai:
        console.print("[red]OpenAI API 키가 없습니다. _local_keys.py 또는 환경변수를 확인하세요.[/red]")
        return

    # ── 세션 선택 ─────────────────────────────────────────────────
    global _current_session_id
    history: list = []
    sessions = _list_sessions(5)

    # 구버전 단일 파일 마이그레이션
    if not sessions and _LEGACY_HISTORY.exists():
        try:
            old = json.loads(_LEGACY_HISTORY.read_text(encoding="utf-8"))
            old_msgs = old.get("history", [])
            if old_msgs:
                _current_session_id = _new_session_id()
                _save_session(_current_session_id, old_msgs, "(이전 버전 기록)")
                _LEGACY_HISTORY.unlink(missing_ok=True)
                sessions = _list_sessions(5)
        except Exception:
            pass

    if sessions:
        last = sessions[0]
        console.print(f"[dim]마지막 대화: {last['last_active']}  {last['turn_count']}턴  \"{last['first_message']}\"[/dim]")
        try:
            choice = Prompt.ask(
                "  [dim]이어서 대화(Enter) / 새 대화(n) / 목록(l)[/dim]",
                default="y"
            ).strip().lower()
        except Exception:
            choice = "y"

        if choice == "l":
            _show_sessions(sessions)
            try:
                idx = Prompt.ask("  불러올 번호 (Enter = 새 대화)", default="0").strip()
                if idx.isdigit() and 1 <= int(idx) <= len(sessions):
                    sid = sessions[int(idx)-1]["sid"]
                    history = _load_session(sid)
                    _current_session_id = sid
                    console.print(f"  [green]{len(history)//2}턴 불러옴[/green]\n")
                else:
                    _current_session_id = _new_session_id()
            except Exception:
                _current_session_id = _new_session_id()
        elif choice in ("", "y", "yes"):
            history = _load_session(last["sid"])
            _current_session_id = last["sid"]
            console.print(f"  [green]{len(history)//2}턴 불러옴[/green]\n")
        else:
            _current_session_id = _new_session_id()
    else:
        _current_session_id = _new_session_id()

    console.print(f"[dim]세션: {_current_session_id}  |  무엇이든 물어보세요. 'exit'로 종료.[/dim]\n")

    while True:
        try:
            user_input = Prompt.ask("[bold magenta]You[/bold magenta]")
        except (KeyboardInterrupt, EOFError):
            break

        if not user_input.strip():
            continue
        if user_input.strip().lower() in ("exit", "quit", "종료"):
            break

        # 특수 명령
        if user_input.strip() == "/models":
            _show_model_status()
            continue
        if user_input.strip() == "/cache":
            _show_cache()
            continue
        if user_input.strip() == "/help":
            _show_help()
            continue
        if user_input.strip() == "/clear":
            history = []
            _current_session_id = _new_session_id()
            console.print(f"[dim]새 세션 시작: {_current_session_id}[/dim]\n")
            continue
        if user_input.strip() == "/sessions":
            _show_sessions(_list_sessions(10))
            try:
                idx = Prompt.ask("  불러올 번호 (Enter = 취소)", default="0").strip()
                if idx.isdigit() and 1 <= int(idx) <= 10:
                    slist = _list_sessions(10)
                    if int(idx) <= len(slist):
                        sid = slist[int(idx)-1]["sid"]
                        history = _load_session(sid)
                        _current_session_id = sid
                        console.print(f"  [green]{len(history)//2}턴 불러옴[/green]\n")
            except Exception:
                pass
            continue

        console.print()

        with console.status("[bold magenta]MAESTRO 처리 중...[/bold magenta]", spinner="dots"):
            start = time.time()
            response = run_agent(user_input, history)
            elapsed = time.time() - start

        console.print()
        console.print(Panel(
            Markdown(response),
            title=f"[bold cyan]MAESTRO[/bold cyan]  [dim]{elapsed:.1f}s[/dim]",
            border_style="cyan", padding=(1, 2)
        ))
        console.print()

        # 히스토리 누적 & 세션 저장 (전체 무제한 로컬 보존)
        first_msg = user_input if len(history) == 0 else ""
        history.append({"role": "user",      "content": user_input})
        history.append({"role": "assistant", "content": response})
        _save_session(_current_session_id, history, first_msg)

        # GPT-4o 컨텍스트는 최근 20턴만 (비용/속도 최적화)
        # history 자체는 전체 유지, run_agent에서 슬라이싱


def _show_sessions(sessions: list):
    from rich.table import Table
    t = Table(title="저장된 세션", border_style="dim")
    t.add_column("번호", style="dim", width=4)
    t.add_column("시작", width=14)
    t.add_column("마지막 활동", width=14)
    t.add_column("턴", width=4)
    t.add_column("첫 메시지")
    for i, s in enumerate(sessions, 1):
        t.add_row(str(i), s["started_at"], s["last_active"],
                  str(s["turn_count"]), s["first_message"][:40])
    console.print(t)


def _show_model_status():
    from rich.table import Table
    t = Table(title="연결된 모델", border_style="dim")
    t.add_column("모델"); t.add_column("상태"); t.add_column("특기")
    t.add_row("GPT-4o",   "[green]OK[/green]" if oai      else "[red]X[/red]", "오케스트레이터, 구조화 출력")
    t.add_row("Claude",   "[green]OK[/green]" if ant      else "[red]X[/red]", "긴 문서, 정밀 분석, 글쓰기")
    t.add_row("DeepSeek", "[green]OK[/green]" if deepseek else "[red]X[/red]", "추론, 수학, 알고리즘")
    t.add_row("Grok",     "[green]OK[/green]" if grok_ai  else "[red]X[/red]", "실시간 정보, 최신 이슈")
    console.print(t)


def _show_cache():
    if not _ORCH:
        console.print("[red]오케스트레이터 비활성[/red]")
        return
    from rich.table import Table
    rows = orch.get_cache_summary()
    t = Table(title="오케스트레이터 캐시", border_style="dim")
    t.add_column("카테고리"); t.add_column("Winner"); t.add_column("Meta"); t.add_column("Self"); t.add_column("Streak")
    for r in rows:
        t.add_row(r["category"], r["winner"],
                  str(r["meta_score"]), str(r["self_score"]), str(r["low_score_streak"]))
    console.print(t)


def _show_help():
    console.print(Panel(
        "[bold]사용 가능한 명령[/bold]\n\n"
        "  /models   : 연결된 LLM 상태\n"
        "  /cache    : 오케스트레이터 캐시 보기\n"
        "  /sessions : 과거 대화 목록 & 불러오기\n"
        "  /clear    : 새 세션 시작\n"
        "  /help     : 이 도움말\n"
        "  exit      : 종료\n\n"
        "[bold]예시 요청[/bold]\n\n"
        "  'C:\\Users\\나\\Desktop\\보고서.pdf 분석해줘'\n"
        "  '스크린샷.png 텍스트 추출해줘'\n"
        "  '데이터.xlsx 핵심 인사이트 뽑아줘'\n"
        "  '삼성전자 최근 마케팅 전략 조사해줘'\n"
        "  '랜딩페이지 만들고 싶어, 어디서 시작해?'\n"
        "  'requirements.txt 만들어줘'",
        border_style="dim"
    ))


if __name__ == "__main__":
    main()
